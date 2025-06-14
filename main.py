import json

from pydantic import BaseModel
from docx import Document
from win32com import client
import os
import tempfile
import pythoncom
from docx.oxml.ns import qn
from typing import Dict, List, Union, Any
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from fastapi.responses import Response
from fastapi.encoders import jsonable_encoder

def extract_header_footer(header_footer) -> List[Dict]:
    """提取页眉/页脚内容"""
    content = []
    for para in header_footer.paragraphs:
        content.append({
            "text": para.text,
            "alignment": str(para.alignment) if para.alignment else None
        })
    return content


def get_highlight_color(run) -> Union[str, None]:
    """获取文本高亮颜色"""
    if run._element.rPr is None:
        return None
    highlight = run._element.rPr.find(qn('w:highlight'))
    return highlight.get(qn('w:val')) if highlight is not None else None


def read_docx_formatting(docx_path):
    doc = Document(docx_path)
    formatting_data = {"paragraphs": [], "tables": []}

    for para in doc.paragraphs:
        run_styles = [{"text": run.text, "bold": run.bold, "italic": run.italic,
                       "underline": run.underline, "font_name": run.font.name,
                       "font_size": run.font.size.pt if run.font.size else None}
                      for run in para.runs]
        formatting_data["paragraphs"].append({
            "text": para.text,
            "style": para.style.name if para.style else None,
            "runs": run_styles
        })

    for table in doc.tables:
        table_data = [[cell.text for cell in row.cells] for row in table.rows]
        formatting_data["tables"].append(table_data)

    sections = []
    for section in doc.sections:
        sect_format = {
            "page_width": section.page_width.inches,  # 纸张宽度
            "page_height": section.page_height.inches,  # 纸张高度
            "left_margin": section.left_margin.inches,  # 纸张左侧页边距
            "right_margin": section.right_margin.inches,  # 纸张右侧页边距
            "top_margin": section.top_margin.inches,  # 纸张上方页边距
            "bottom_margin": section.bottom_margin.inches,  # 纸张下方页边距
            "header_distance": section.header_distance.inches,  # 页眉位置
            "footer_distance": section.footer_distance.inches,  # 页脚位置
            "gutter": section.gutter.inches,  # 装订线
            "orientation": "landscape" if section.orientation else "portrait",  # 纸张方向
            "header": extract_header_footer(section.header),  # 页眉内容
            "footer": extract_header_footer(section.footer),  # 页脚内容
            # 注释：文档可以包含多个具有不同页面设置的节
        }
        sections.append(sect_format)
    formatting_data["sections"] = sections

    return formatting_data


app = FastAPI(title="Word Format Extractor API")


def convert_doc_to_docx(input_path, output_path):
    pythoncom.CoInitialize()
    word = client.Dispatch("Word.Application")
    doc = word.Documents.Open(input_path)
    doc.SaveAs(output_path, FileFormat=16)  # 16 = wdFormatDocumentDefault (.docx)
    doc.Close()
    word.Quit()
    pythoncom.CoUninitialize()


@app.post("/extract-format")
async def extract_format(file: UploadFile = File(...)):
    if not file.filename.endswith((".doc", ".docx")):
        raise HTTPException(status_code=400, detail="请上传 .doc 或 .docx 文件")

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, file.filename)
            with open(input_path, "wb") as f:
                f.write(await file.read())

            if input_path.endswith(".doc"):
                docx_path = os.path.join(tmpdir, "converted.docx")
                convert_doc_to_docx(input_path, docx_path)
            else:
                docx_path = input_path

            formatting_info = read_docx_formatting(docx_path)
            return JSONResponse(content=jsonable_encoder(formatting_info))
            # return Response(
            #     content=json.dumps(
            #         formatting_info,
            #         ensure_ascii=False,
            #         indent=4,
            #     ),
            #     media_type="application/json",
            # )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理文件时出错: {str(e)}")

# 运行接口：uvicorn main:app --reload
# 测试接口：curl -X POST "http://127.0.0.1:8000/extract-format" -F "file=@E:\lyzzzz\formatDetection\code\files\temp_converted.docx"



